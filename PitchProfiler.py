import pandas as pd
from pybaseball import statcast_pitcher
from pybaseball import playerid_lookup
import matplotlib.pylab as plt
import matplotlib.gridspec as gridspec
import math as math
from pandas.compat import BytesIO
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
import os

plt.style.use('seaborn-paper')


# used for debugging
pd.set_option('display.max_columns', None)
pd.set_option('display.expand_frame_repr', False)
pd.set_option('max_colwidth', -1)


def get_number(last, first):
    playerTable = playerid_lookup(last, first)
    playerTable = playerTable.loc[playerTable['mlb_played_last'].isin([2019])]
    playerTable.index = range(len(playerTable['mlb_played_last']))
    number = playerTable['key_mlbam']
    number = number[0]
    return number


def import_data(number, start, end):
    data = statcast_pitcher(start_dt=start, end_dt=end,
                            player_id=number)
    data = data[['pitch_type', 'release_speed', 'release_pos_x', 'release_pos_z',
                 'pfx_x', 'pfx_z', 'release_spin_rate', 'plate_x', 'plate_z',
                 'estimated_woba_using_speedangle', 'woba_value', 'description',
                 'launch_speed_angle', 'launch_angle', 'launch_speed', 'bb_type',
                 'effective_speed', 'vx0', 'vy0', 'vz0', 'ax', 'ay', 'az',
                 'release_extension']]
    data.index = range(len(data['pitch_type']))
    return data


def nathan_calculations(pitches):
    # constants
    g_fts = 32.174
    R_ball = .121
    mass = 5.125
    circ = 9.125
    temp = 72
    humidity = 50
    pressure = 29.92
    temp_c = (5/9)*(temp-32)
    pressure_mm = (pressure * 1000) / 39.37
    svp = 4.5841 * math.exp((18.687 - temp_c/234.5) * temp_c/(257.14 + temp_c))
    rho = (1.2929 * (273 / (temp_c + 273)) * (pressure_mm - .3783 *
                                              humidity * svp / 100) / 760) * .06261
    const = 0.07182 * rho * (5.125 / mass) * (circ / 9.125)**2

    # add row to put calculations in
    pitches['InducedHorzBreak'] = np.nan
    pitches['InducedVertBreak'] = np.nan
    pitches['Tilt'] = np.nan
    pitches['SpinEff'] = np.nan
    lol = 0
    for i in range(len(pitches.pitch_type)):
        # x0 = -1 * pitches.x0.iloc[i]
        v0 = pitches.release_speed.iloc[i]
        vx0 = pitches.vx0.iloc[i]
        ax = pitches.ax.iloc[i]
        vy0 = pitches.vy0.iloc[i]
        ay = pitches.ay.iloc[i]
        vz0 = pitches.vz0.iloc[i]
        az = pitches.az.iloc[i]
        pfx_x = pitches.pfx_x.iloc[i]
        pfx_z = pitches.pfx_z.iloc[i]
        plate_x = pitches.plate_x.iloc[i]
        plate_z = pitches.plate_z.iloc[i]
        release_x = pitches.release_pos_x.iloc[i]
        release_y = 60.5-pitches.release_extension.iloc[i]
        release_z = pitches.release_pos_z.iloc[i]
        spin_rate = pitches.release_spin_rate.iloc[i]

        # time between release and y0 measurement
        t_back_to_release = (-vy0-math.sqrt(vy0**2-2*ay*(50-release_y)))/ay

        # adjust velocity at y0 to be at release
        vx_r = vx0+ax*t_back_to_release
        vy_r = vy0+ay*t_back_to_release
        vz_r = vz0+az*t_back_to_release
        dv0 = v0 - math.sqrt(vx_r**2 + vy_r**2 + vz_r**2)/1.467

        # calculate pitch time also know as tf in Template
        t_c = (-vy_r - math.sqrt(vy_r**2 - 2*ay*(release_y - 17/12))) / ay

        # calcualte x and z movement
        calc_x_mvt = (plate_x-release_x-(vx_r/vy_r)*(17/12-release_y))
        calc_z_mvt = (plate_z-release_z-(vz_r/vy_r)*(17/12-release_y))+0.5*g_fts*t_c**2

        # average velocity
        vx_bar = (2 * vx_r + ax * t_c) / 2
        vy_bar = (2 * vy_r + ay * t_c) / 2
        vz_bar = (2 * vz_r + az * t_c) / 2
        v_bar = math.sqrt(vx_bar**2 + vy_bar**2 + vz_bar**2)

        # drag acceleration
        adrag = -(ax * vx_bar + ay * vy_bar + (az + g_fts) * vz_bar)/v_bar

        # magnus acceleration
        amagx = ax + adrag * vx_bar/v_bar
        amagy = ay + adrag * vy_bar/v_bar
        amagz = az + adrag * vz_bar/v_bar + g_fts
        amag = math.sqrt(amagx**2 + amagy**2 + amagz**2)

        # movement components
        mx = .5 * amagx * (t_c**2)*12
        mz = .5 * amagz * (t_c**2)*12

        # drag/lift coefficients may need work
        Cd = adrag / (v_bar**2 * const)
        Cl = amag / (v_bar**2 * const)

        s = 0.4*Cl/(1-2.32*Cl)
        spin_t = 78.92*s*v_bar

        '''
        # for debugging purposes
        spin_tx = spin_t*(vy_bar*amagz-vz_bar*amagy)/(amag*v_bar)
        spin_ty = spin_t*(vz_bar*amagx-vx_bar*amagz)/(amag*v_bar)
        spin_tz = spin_t*(vx_bar*amagy-vy_bar*amagx)/(amag*v_bar)
        spin_check = math.sqrt(spin_tx**2+spin_ty**2+spin_tz**2)-spin_t
        '''
        # calc spin direction
        phi = 0
        if(amagz > 0):
            phi = math.atan2(amagz, -amagx) * 180/math.pi
        else:
            phi = 360+math.atan2(amagz, -amagx)*180/math.pi
        dec_time = 3-(1/30)*phi
        if(dec_time <= 0):
            dec_time += 12

        # calc spin eff
        spin_eff = spin_t/spin_rate
        pd.set_option('mode.chained_assignment', None)
        pitches.InducedHorzBreak.iloc[i] = -calc_x_mvt
        pitches.InducedVertBreak.iloc[i] = calc_z_mvt
        pitches.Tilt.iloc[i] = dec_time
        pitches.SpinEff.iloc[i] = spin_eff

    return pitches


def get_pitch_types(data):
    pitchcounts = data['pitch_type'].value_counts(dropna=True)
    pitch_types = []
    for i in range(len(pitchcounts)):
        is_pitch = pitchcounts.index[i]
        pitch_types.append(is_pitch)
    return pitch_types


def color_picker(pitch_type):
    color = ''
    if(pitch_type == 'FF'):
        color = '#8C1C13'
    elif(pitch_type == 'FT'):
        color = '#CC5803'
    elif(pitch_type == 'SI'):
        color = '#FE7F2D'
    elif(pitch_type == 'FC'):
        color = '#E08E45'
    elif(pitch_type == 'FS'):
        color = '#F3CA4C'
    elif(pitch_type == 'SL'):
        color = '#274060'
    elif(pitch_type == 'CU'):
        color = '#4EA5D9'
    elif(pitch_type == 'KC'):
        color = '#5BC0EE'
    elif(pitch_type == 'CH'):
        color = '#1446A0'
    elif(pitch_type == 'KN'):
        color = '#712F79'
    elif(pitch_type == 'FO'):
        color = '#03B5AA'
    elif(pitch_type == 'EP'):
        color = '#DBFE97'
    elif(pitch_type == 'SC'):
        color = '#5B9279'
    else:
        color = 'black'
    return color


def plot_release_movement(data):
    pitch_types = get_pitch_types(data)
    gs = gridspec.GridSpec(2, 2)
    fig = plt.figure(figsize=(6, 2.5))
    ax0 = plt.subplot(gs[:, 0])  # release
    ax2 = plt.subplot(gs[:, 1])  # break
    for i in range(len(pitch_types)):
        is_pitch = data['pitch_type'] == pitch_types[i]
        selected_data = data[is_pitch]
        pitch_count = selected_data['pitch_type'].count()
        total_count = data['pitch_type'].count()
        percent = round(pitch_count/total_count*100, 1)
        label = pitch_types[i]

        if(percent < 1.5):
            continue
        else:
            color = color_picker(label)
            if(label == 'PO' or label == 'IB' or label == 'AB' or label == 'UN'):
                continue
            else:
                ax0.scatter(-selected_data['release_pos_x'], selected_data['release_pos_z'],
                            label=label, s=20, alpha=0.5, c=color)
                ax2.scatter(12*selected_data['InducedHorzBreak'], 12*selected_data['InducedVertBreak'],
                            label=label, s=20, alpha=0.5, c=color)

    ax0.set_xlim(-data['release_pos_x'].mean()-1.5, -data['release_pos_x'].mean()+1.5)
    ax0.set_ylim(data['release_pos_z'].mean()-1.5, data['release_pos_z'].mean()+1.5)
    ax0.set_title('Release Position (Pitchers View)')
    ax0.set_xlabel('Horizontal Release Point')
    ax0.set_ylabel('Vertical Release Point')
    ax2.set_xlim(-30, 30)
    ax2.set_ylim(-30, 30)
    ax2.axhline(y=0, color='k')
    ax2.axvline(x=0, color='k')
    ax2.set_title('Pitch Movement (Pitcher\'s View)')
    ax2.set_xlabel('Horizontal Movement')
    ax2.set_ylabel('Vertical Movement')
    ax0.legend(prop={'size': 7})
    ax2.legend(prop={'size': 7})
    fig.tight_layout()
    fig.subplots_adjust(top=.92, wspace=0.35, bottom=.18, left=.11, right=.98)
    memfile = BytesIO()
    plt.savefig(memfile)
    # plt.show() for debug
    return memfile


def plot_location(data):
    pitch_types = get_pitch_types(data)
    gs = gridspec.GridSpec(1, 7)
    fig = plt.figure(figsize=(6, 1.5))
    ax0 = plt.subplot(gs[:, 0])  # FF
    ax1 = plt.subplot(gs[:, 1])  # FT/SI
    ax2 = plt.subplot(gs[:, 2])  # FC
    ax3 = plt.subplot(gs[:, 3])  # SL
    ax4 = plt.subplot(gs[:, 4])  # CU/KC
    ax5 = plt.subplot(gs[:, 5])  # CH
    ax6 = plt.subplot(gs[:, 6])  # FS

    for i in range(len(pitch_types)):
        is_pitch = data['pitch_type'] == pitch_types[i]
        selected_data = data[is_pitch]
        label = pitch_types[i]
        pitch_count = selected_data['pitch_type'].count()
        total_count = data['pitch_type'].count()
        percent = round(pitch_count/total_count*100, 1)
        label = pitch_types[i]
        if(percent < 1.5):
            continue
        else:
            sz_x = [.79, .79, -.79, -.79, .79]
            sz_z = [3.5, 1.5, 1.5, 3.5, 3.5]
            xedges, zedges = np.linspace(-2, 2, 20), np.linspace(-0.5, 4.5, 20)
            x = -1*selected_data['plate_x']
            z = selected_data['plate_z']
            hist, xedges, yedges = np.histogram2d(x, z, (xedges, zedges))
            xidx = np.clip(np.digitize(x, xedges), 0, hist.shape[0]-1)
            zidx = np.clip(np.digitize(z, zedges), 0, hist.shape[1]-1)
            c = hist[xidx, zidx]
            if(label == 'FF'):
                ax0.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax0.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'FT' or label == 'SI'):
                ax1.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax1.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'FC'):
                ax2.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax2.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'SL'):
                ax3.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax3.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'CU' or label == 'KC'):
                ax4.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax4.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'CH'):
                ax5.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax5.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'FS'):
                ax6.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax6.plot(sz_x, sz_z, color='black', lw=0.5)

    ax0.set_xlim(-2, 2)
    ax0.set_ylim(-0.5, 4.5)
    ax0.axis('off')
    ax0.set_title('FF', fontsize=10)
    ax1.set_xlim(-2, 2)
    ax1.set_ylim(-0.5, 4.5)
    ax1.axis('off')
    ax1.set_title('FT/SI', fontsize=10)
    ax2.set_xlim(-2, 2)
    ax2.set_ylim(-0.5, 4.5)
    ax2.axis('off')
    ax2.set_title('FC', fontsize=10)
    ax3.set_xlim(-2, 2)
    ax3.set_ylim(-0.5, 4.5)
    ax3.axis('off')
    ax3.set_title('SL', fontsize=10)
    ax4.set_xlim(-2, 2)
    ax4.set_ylim(-0.5, 4.5)
    ax4.axis('off')
    ax4.set_title('CU/KC', fontsize=10)
    ax5.set_xlim(-2, 2)
    ax5.set_ylim(-0.5, 4.5)
    ax5.axis('off')
    ax5.set_title('CH', fontsize=10)
    ax6.set_xlim(-2, 2)
    ax6.set_ylim(-0.5, 4.5)
    ax6.axis('off')
    ax6.set_title('FS', fontsize=10)
    fig.suptitle('Pitch Locations (Pitcher\'s View)')
    fig.subplots_adjust(top=.60, wspace=0.0, bottom=.0, left=.00, right=1)
    memfile = BytesIO()
    plt.savefig(memfile)
    # plt.show() for debug
    return memfile


# rounds to nearest 15 minutes
def time_round(x, base=15):
    return base * round(x/base)


def convert_to_time(dec_time):
    if(math.isnan(dec_time)):
        time = ''
    else:
        hours = int(dec_time)
        minutes = int((dec_time*60) % 60)
        minutes = time_round(minutes)
        if(minutes == 60):
            hours += 1
            minutes = 0
        if(hours == 0):
            hours = 12
        elif(hours > 12):
            hours = hours - 12
        minutestr = str(minutes)
        if(len(minutestr) < 2):
            minutestr = '0' + minutestr
        time = str(hours) + ':' + minutestr
    return time


def transform_data(data):
    pitch_types = get_pitch_types(data)
    pitches = []
    battedball = []
    for i in range(len(pitch_types)):
        is_pitch = data['pitch_type'] == pitch_types[i]
        selected_data = data[is_pitch]
        count = selected_data['pitch_type'].count()
        if (count > 0):
            label = pitch_types[i]
            # pitch info stuff
            swm = selected_data[selected_data['description'] == 'swinging_strike'].count()[
                'description']
            desiredOutcome = ['swinging_strike', 'swinging_strike_blocked', 'foul',
                              'foul_tip', 'hit_into_play', 'hit_into_play_no_out', 'hit_into_play_score']
            total_swings = selected_data.loc[selected_data['description'].isin(desiredOutcome)].count()[
                'description']
            percentage_used = round(
                (count/data['pitch_type'].count()) * 100, 1)
            avgVelo = round(selected_data['release_speed'].dropna().mean(), 1)
            avgSpinRate = round(selected_data['release_spin_rate'].dropna().mean(), 0)
            avgHorzBreak = round(12*selected_data['InducedHorzBreak'].dropna().mean(), 1)
            avgVertBreak = round(12*selected_data['InducedVertBreak'].dropna().mean(), 1)
            whiff_rate = round(swm/total_swings*100, 1)
            #bauer_units = (round(avgSpinRate/avgVelo, 0))
            tilt = convert_to_time(selected_data['Tilt'].dropna().mean())
            spin_eff = round(selected_data['SpinEff'].dropna().mean()*100, 1)

            # batted ball stuff
            bbe = selected_data['launch_speed_angle'].dropna().count()
            weak = round(selected_data[selected_data['launch_speed_angle'] == 1].count()[
                'launch_speed_angle']/bbe*100, 1)
            topped = round(selected_data[selected_data['launch_speed_angle'] == 2].count()[
                'launch_speed_angle']/bbe*100, 1)
            under = round(selected_data[selected_data['launch_speed_angle'] == 3].count()[
                'launch_speed_angle']/bbe*100, 1)
            flare = round(selected_data[selected_data['launch_speed_angle'] == 4].count()[
                'launch_speed_angle']/bbe*100, 1)
            solid = round(selected_data[selected_data['launch_speed_angle'] == 5].count()[
                'launch_speed_angle']/bbe*100, 1)
            barrels = round(selected_data[selected_data['launch_speed_angle'] == 6].count()[
                'launch_speed_angle']/bbe*100, 1)
            hardhit = round(selected_data[selected_data['launch_speed'] >= 95].count()[
                'launch_speed_angle']/bbe*100, 1)

            pitch = [label, percentage_used, avgVelo, avgSpinRate,
                     avgHorzBreak, avgVertBreak, tilt, spin_eff, whiff_rate]
            bbs = [label, percentage_used, weak, topped, under, flare, solid,
                   barrels, hardhit]
            pitches.append(pitch)
            battedball.append(bbs)
        else:
            continue
    pitches = pd.DataFrame(pitches, columns=['Pitch Type', '% Thrown', 'Velocity',
                                             'Spin Rate', 'Horizontal Break',
                                             'Vertical Break', 'Tilt', 'Spin Eff.',
                                             'Whiff Rate'])
    battedball = pd.DataFrame(battedball, columns=['Pitch Type', '% Thrown',
                                                   'Weak%', 'Topped%', 'Under%',
                                                   'Burner%', 'Solid%',
                                                   'Barrel%', 'HardHit%'])
    pitches = pitches[pitches['% Thrown'] >= 1.5]
    battedball = battedball[battedball['% Thrown'] >= 1.5]
    pitches = pitches.sort_values(by=['% Thrown'], ascending=False)
    battedball = battedball.sort_values(by=['% Thrown'], ascending=False)

    return pitches, battedball


def generate_profile(fname, lname, date1, date2, moves, locs, pitches, battedballs, count):
    document = Document()

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titlestr = 'Pitch Profile for ' + fname + ' ' + lname
    datestr = '\n' + date1 + ' to ' + date2
    run0 = paragraph.add_run(titlestr)
    run1 = paragraph.add_run(datestr)
    font0 = run0.font
    font0.size = Pt(24)
    font0.bold = True
    font0.name = 'Calibri'
    font1 = run1.font
    font1.size = Pt(12)
    font1.name = 'Calibri'

    move_paragraph = document.add_paragraph()
    paragraph_format = move_paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = move_paragraph.add_run()
    my_image = run.add_picture(moves)

    loc_paragraph = document.add_paragraph()
    paragraph_format = loc_paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = loc_paragraph.add_run()
    my_image = run.add_picture(locs)

    # First row are table headers
    pitchTable = document.add_table(
        pitches.shape[0]+1, pitches.shape[1], style='MediumList2')

    # add the header rows.
    for j in range(pitches.shape[-1]):
        pitchTable.cell(0, j).text = pitches.columns[j]

    # add the rest of the data frame
    for i in range(pitches.shape[0]):
        for j in range(pitches.shape[-1]):
            pitchTable.cell(i+1, j).text = str(pitches.values[i, j])

    break_paragraph = document.add_paragraph('')

    # First row are table headers
    bbTable = document.add_table(
        battedballs.shape[0]+1, battedballs.shape[1], style='MediumList2')

    # add the header rows.
    for j in range(battedballs.shape[-1]):
        bbTable.cell(0, j).text = battedballs.columns[j]

    # add the rest of the data frame
    for i in range(battedballs.shape[0]):
        for j in range(battedballs.shape[-1]):
            bbTable.cell(i+1, j).text = str(battedballs.values[i, j])

    document.save(fname + lname + date2 + '.docx')
    moves.close()
    locs.close()


def main():
    while (1 == 1):
        fname = input('Enter First Name: ')
        lname = input('Enter Last Name: ')
        date1 = input('Enter Start of Date Range (YYYY-MM-DD): ')
        date2 = input('Enter End of Date Range (YYYY-MM-DD): ')
        data = import_data(get_number(lname, fname), date1, date2)
        data = nathan_calculations(data)
        pitchdata, battedballdata = transform_data(data)
        releasemovement = plot_release_movement(data)
        locations = plot_location(data)
        pitches = get_pitch_types(data)
        count = 0
        for i in range(len(pitches)):
            count += 1

        generate_profile(fname, lname, date1, date2, releasemovement,
                         locations, pitchdata, battedballdata, count)

        unused_variable = os.system("cls")
        print('Pitch Profile for: ' + fname + lname + date2 + '.docx Created!')
        dec = input('Enter New Pitcher (Y/N)?: ').upper()
        if(dec == 'N'):
            break
        else:
            continue


if __name__ == '__main__':
    main()
