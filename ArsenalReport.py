import pandas as pd
from pybaseball import statcast_pitcher
from pybaseball import playerid_lookup
import matplotlib.pylab as plt
import matplotlib.gridspec as gridspec
import math as m
from pandas.compat import BytesIO
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

plt.style.use('seaborn-paper')

'''
# used for debugging
pd.set_option('display.max_columns', None)
pd.set_option('display.expand_frame_repr', False)
pd.set_option('max_colwidth', -1)
'''


def getNumber(last, first):
    playerTable = playerid_lookup(last, first)
    playerTable = playerTable.loc[playerTable['mlb_played_last'].isin([2019])]
    playerTable.index = range(len(playerTable['mlb_played_last']))
    number = playerTable['key_mlbam']
    number = number[0]
    return number


def dataGrab(number, start, end):
    data = statcast_pitcher(start_dt=start, end_dt=end,
                            player_id=number)
    data = data[['pitch_type', 'release_speed', 'release_pos_x', 'release_pos_z',
                 'pfx_x', 'pfx_z', 'release_spin_rate',
                 'estimated_woba_using_speedangle', 'woba_value', 'description']]
    data.index = range(len(data['pitch_type']))
    return data


def getPitchTypes(data):
    pitchcounts = data['pitch_type'].value_counts(dropna=True)
    pitch_types = []
    for i in range(len(pitchcounts)):
        is_pitch = pitchcounts.index[i]
        pitch_types.append(is_pitch)
    return pitch_types


def colorPicker(pitch_type):
    color = ''
    if(pitch_type == 'FF'):
        color = '#8C1C13'
    elif(pitch_type == 'FT'):
        color = '#CC5803'
    elif(pitch_type == 'FT'):
        color = '#FE7F2D'
    elif(pitch_type == 'FC'):
        color = '#E08E45'
    elif(pitch_type == 'FS'):
        color = '#F3CA4C'
    elif(pitch_type == 'SL'):
        color = '#274060'
    elif(pitch_type == 'CU'):
        color = '#4EA5D9'
    elif(pitch_type == 'KC'):
        color = '#5BC0EE'
    elif(pitch_type == 'CH'):
        color = '#1446A0'
    elif(pitch_type == 'KN'):
        color = '#712F79'
    elif(pitch_type == 'FO'):
        color = '#03B5AA'
    elif(pitch_type == 'EP'):
        color = '#DBFE97'
    elif(pitch_type == 'SC'):
        color = '#5B9279'
    else:
        color = 'black'
    return color


def plotData(data):
    pitch_types = getPitchTypes(data)
    gs = gridspec.GridSpec(2, 2)
    fig = plt.figure(figsize=(6, 3))
    ax0 = plt.subplot(gs[:, 0])  # release
    ax2 = plt.subplot(gs[:, 1])  # break
    for i in range(len(pitch_types)):
        is_pitch = data['pitch_type'] == pitch_types[i]
        selected_data = data[is_pitch]
        label = pitch_types[i]
        color = colorPicker(label)
        if(label == 'PO' or label == 'IB' or label == 'AB' or label == 'UN'):
            continue
        else:
            ax0.scatter(selected_data['release_pos_x'], selected_data['release_pos_z'],
                        label=label, s=20, alpha=0.5, c=color)
            ax2.scatter(12*selected_data['pfx_x'], 12*selected_data['pfx_z'],
                        label=label, s=20, alpha=0.5, c=color)
    ax0.set_xlim(data['release_pos_x'].mean()-1.5, data['release_pos_x'].mean()+1.5)
    ax0.set_ylim(data['release_pos_z'].mean()-1.5, data['release_pos_z'].mean()+1.5)
    ax0.set_title('Release Position')
    ax0.set_xlabel('Horizontal Release Point')
    ax0.set_ylabel('Vertical Release Point')
    ax2.set_xlim(-30, 30)
    ax2.set_ylim(-30, 30)
    ax2.axhline(y=0, color='k')
    ax2.axvline(x=0, color='k')
    ax2.set_title('Pitch Movement')
    ax2.set_xlabel('Horizontal Movement')
    ax2.set_ylabel('Vertical Movement')
    ax0.legend(prop={'size': 7})
    ax2.legend(prop={'size': 7})
    fig.tight_layout()
    fig.subplots_adjust(top=.92, wspace=0.35, bottom=.18, left=.11, right=.98)
    memfile = BytesIO()
    plt.savefig(memfile)
    # plt.show() for debug
    return memfile


def getData(data):
    pitch_types = getPitchTypes(data)
    pitches = []
    for i in range(len(pitch_types)):
        is_pitch = data['pitch_type'] == pitch_types[i]
        selected_data = data[is_pitch]
        label = pitch_types[i]
        count = selected_data['pitch_type'].count()
        swstr = selected_data[selected_data['description'] == 'swinging_strike'].count()[
            'description']
        percentage_used = round(
            (count/data['pitch_type'].count()) * 100, 1)

        avgVelo = round(selected_data['release_speed'].dropna().mean(), 1)
        avgSpinRate = round(selected_data['release_spin_rate'].dropna().mean(), 0)
        avgHorzBreak = round(12*selected_data['pfx_x'].dropna().mean(), 1)
        avgVertBreak = round(12*selected_data['pfx_z'].dropna().mean(), 1)
        # avgBreakDir = round(m.degrees(m.atan2(avgHorzBreak, avgVertBreak)), 0)
        estwOBA = round(selected_data['estimated_woba_using_speedangle'].dropna().mean(), 3)
        wOBA = round(selected_data['woba_value'].dropna().mean(), 3)
        swstrperc = round(swstr/count*100, 1)

        pitch = [label, percentage_used, avgVelo, avgSpinRate, avgHorzBreak,
                 avgVertBreak, estwOBA, wOBA, swstrperc]
        pitches.append(pitch)
    pitches = pd.DataFrame(pitches, columns=['Pitch Type', '% Thrown',
                                             'Velocity (mph)', 'Spin Rate (rpm)',
                                             'Horizontal Break (in)', 'Vertical Break (in)',
                                             'xwOBA', 'wOBA', 'SwStr%'])
    pitches = pitches[pitches['% Thrown'] >= 1]
    pitches = pitches.sort_values(by=['% Thrown'], ascending=False)

    return pitches


def GenerateReport(fname, lname, date1, date2, memfile, reportData, count):
    document = Document()

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titlestr = 'Arsenal Report for ' + fname + ' ' + lname
    datestr = '\n' + date1 + ' to ' + date2
    run0 = paragraph.add_run(titlestr)
    run1 = paragraph.add_run(datestr)
    font0 = run0.font
    font0.size = Pt(24)
    font0.bold = True
    font0.name = 'Calibri'
    font1 = run1.font
    font1.size = Pt(12)
    font1.name = 'Calibri'

    next_paragraph = document.add_paragraph()
    paragraph_format = next_paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = next_paragraph.add_run()
    my_image = run.add_picture(memfile)

    last_paragraph = document.add_paragraph()
    headerstr = 'Aresenal'
    run = last_paragraph.add_run(headerstr)
    font2 = run.font
    font2.size = Pt(18)
    font2.bold = True
    font2.name = 'Calibri'

    # First row are table headers
    t = document.add_table(
        reportData.shape[0]+1, reportData.shape[1], style='MediumList2')

    # add the header rows.
    for j in range(reportData.shape[-1]):
        t.cell(0, j).text = reportData.columns[j]

    # add the rest of the data frame
    for i in range(reportData.shape[0]):
        for j in range(reportData.shape[-1]):
            t.cell(i+1, j).text = str(reportData.values[i, j])

    pbreak = document.add_paragraph('\n')

    pitchtypes = document.add_paragraph(style='List Bullet')
    ptstr = 'Pitch Type - FF: Four-Seam Fastball, FT: Two-Seam Fastball, SI: Sinkker, FC: Cutter, FS: Splitter, SL: Slider, CU: Curveball, KC: Knuckle-curve, CH: Changeup, FO: Forkball, SC: Screwball, KN: Knuckleball, EP: Eephus'
    run = pitchtypes.add_run(ptstr)
    font3 = run.font
    font3.size = Pt(8)
    font3.name = 'Calibri'

    velo = document.add_paragraph(style='List Bullet')
    velostr = 'Velocity - recorded in miles per hour at release.'
    run = velo.add_run(velostr)
    font3 = run.font
    font3.size = Pt(8)
    font3.name = 'Calibri'

    spinrate = document.add_paragraph(style='List Bullet')
    spinstr = 'Spin Rate - recorded in revolutions per minute at release.'
    run = spinrate.add_run(spinstr)
    font3 = run.font
    font3.size = Pt(8)
    font3.name = 'Calibri'

    hbreak = document.add_paragraph(style='List Bullet')
    hbreakstr = 'Horizontal Break - horizontal movement, in inches, of the pitch between the release point and home plate, as compared to a theoretical pitch thrown at the same speed with no spin-induced movement.'
    run = hbreak.add_run(hbreakstr)
    font3 = run.font
    font3.size = Pt(8)
    font3.name = 'Calibri'

    vbreak = document.add_paragraph(style='List Bullet')
    vbreakstr = 'Vertical Break - vertical movement, in inches, of the pitch between the release point and home plate, as compared to a theoretical pitch thrown at the same speed with no spin-induced movement.'
    run = vbreak.add_run(vbreakstr)
    font3 = run.font
    font3.size = Pt(8)
    font3.name = 'Calibri'

    xwOBA = document.add_paragraph(style='List Bullet')
    xwOBAstr = 'xwOBA -  formulated using exit velocity, launch angle and, on certain types of batted balls, Sprint Speed.'
    run = xwOBA.add_run(xwOBAstr)
    font3 = run.font
    font3.size = Pt(8)
    font3.name = 'Calibri'

    wOBA = document.add_paragraph(style='List Bullet')
    wOBAstr = 'wOBA -  a statistic, based on linear weights, designed to measure a player\'s overall offensive contributions per plate appearance.'
    run = wOBA.add_run(wOBAstr)
    font3 = run.font
    font3.size = Pt(8)
    font3.name = 'Calibri'

    swstr = document.add_paragraph(style='List Bullet')
    swstrstr = 'SwStr% - measures swing and miss rate based on pitch type. EX: (FF Swings and Misses)/(Total # of FF Thrown) * 100.'
    run = swstr.add_run(swstrstr)
    font3 = run.font
    font3.size = Pt(8)
    font3.name = 'Calibri'

    document.save(fname + lname + date2 + '.docx')
    memfile.close()


def main():
    while (1 == 1):
        fname = input('Enter First Name: ')
        lname = input('Enter Last Name: ')
        date1 = input('Enter Start of Date Range (YYYY-MM-DD): ')
        date2 = input('Enter End of Date Range (YYYY-MM-DD): ')
        data = dataGrab(getNumber(lname, fname), date1, date2)
        reportData = getData(data)
        memfile = plotData(data)
        pitches = getPitchTypes(data)
        count = 0
        for i in range(len(pitches)):
            count += 1
        GenerateReport(fname, lname, date1, date2, memfile, reportData, count)
        unused_variable = os.system("cls")
        print('Arsenal Report: ' + fname + lname + date2 + '.docx Created!')

        dec = input('Enter New Pitcher (Y/N)?: ').upper()
        if(dec == 'N'):
            break
        else:
            continue


if __name__ == '__main__':
    main()
